import os
import nltk
import ssl
import time
import re
from langchain_community.document_loaders import (
    PyPDFLoader, 
    Docx2txtLoader, 
    TextLoader, 
    UnstructuredMarkdownLoader,
    UnstructuredRTFLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
#from langchain_ollama import ChatOllama
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from dotenv import load_dotenv
import requests 
import json
from .rag_config import load_rag_config
#from .databridge_services import run_data_bridge_agent
try:
    from .databridge_services.langgraph_agent import run_data_bridge_agent
except Exception as e:
    run_data_bridge_agent = None
    print(f"⚠️ Optional databridge import failed: {e}")
from app.services.stream_manager import stream_manager

load_dotenv()

class LLMService:
    """Service for LLM provider interactions and RAG pipeline"""
    
    def __init__(self):
        self.rag_config = load_rag_config()

        # Initialize NLTK
        try:
            _create_unverified_https_context = ssl._create_unverified_context
        except AttributeError:
            pass
        else:
            ssl._create_default_https_context = _create_unverified_https_context

        # Initialize NLTK - forcing the download of both required pieces
        try:
            nltk.download('punkt')
            nltk.download('punkt_tab')
        except Exception as e:
            print(f"NLTK Download Warning: {e}")
        # --- FIX ENDS HERE ---
    
        
        # Initialize Embeddings Model
        self.embeddings_model = HuggingFaceEmbeddings(
            model_name=self.rag_config["embedding"]["model"]
        )

        # Permanent Docker Config
        self.qdrant_url = self.rag_config["vector_store"]["url"]
        self.collection_name = self.rag_config["vector_store"]["collection_name"]

        #self.llm = ChatOpenAI(
        #    model="gpt-4o-mini",
        #    temperature=0,
        #    openai_api_key=os.getenv("OPENAI_API_KEY")
        #)

        self.openai_key = os.getenv("OPENAI_API_KEY")
        
        self.models = {
                "gpt-4o-mini": ChatOpenAI(
                    model="gpt-4o-mini",
                    temperature=self.rag_config["generation"]["temperature"],
                    openai_api_key=self.openai_key
                ),
                "gpt-4o": ChatOpenAI(
                    model="gpt-4o",
                    temperature=self.rag_config["generation"]["temperature"],
                    openai_api_key=self.openai_key
                )
            }
        
        # Keep this as a default fallback pointer if needed
        self.llm = self.models["gpt-4o-mini"]

        # Permanent Docker Config for Ollama (Direct API)
        self.ollama_config = {
            "type": "ollama",
            "url": "http://ollama:11434/api/generate",
            "model": "phi3:mini",
            "max_tokens": self.rag_config["generation"]["max_tokens"],
            "temperature": self.rag_config["generation"]["temperature"],
            "timeout": self.rag_config["generation"]["timeout_seconds"],
            "num_ctx": self.rag_config["generation"]["ollama_context_window"]
        }

    #     BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    #     schema_path = os.path.join(BASE_DIR, 'databridge_services', 'sap_schema_with_sap_comments.json')
        
    #     try:
    #         with open(schema_path, 'r') as f:
    #             self.schema_data = json.load(f)
            
    #         # Pre-compile human-readable string version for LLM context fallback
    #         tables_dict = self.schema_data.get("tables", self.schema_data)
            
    #         schema_lines = []
    #         for table_name, table_info in tables_dict.items():
    #             cols_dict = table_info.get("columns", {})
    #             cols_list = list(cols_dict.keys())
    #             schema_lines.append(f"Table '{table_name}': columns = {', '.join(cols_list)}")
            
    #         self.schema_str = "\n".join(schema_lines)
    #         print("✅ Schema loaded successfully into LLMService memory")
    #     except Exception as e:
    #         print(f"⚠️ Schema loading failed: {e}")
    #         self.schema_data = {}
    #         self.schema_str = "No schema available"


    #     # ============================================
    #     # PREBUILD LOOKUP KEYWORDS FOR FAST HEURISTICS
    #     # ============================================
    #     self.schema_keywords = self.build_schema_keywords()
    #     print(f"✅ Loaded {len(self.schema_keywords)} quick-lookup schema tokens")

    #     self.sql_patterns = {
    #         "list", "show", "count", "total", "find", "fetch", 
    #         "display", "records", "entries", "all", "how many", 
    #         "top", "highest", "lowest", "average", "sum"
    #     }

    # # --- Fast Lookup Helper Methods ---
    
    # def build_schema_keywords(self):
    #     keywords = set()
    #     # Safe handling for the "tables" wrapper key in your schema JSON layout
    #     tables_dict = self.schema_data.get("tables", self.schema_data)
        
    #     for table_name, table_info in tables_dict.items():
    #         keywords.add(str(table_name).lower())
            
    #         # Extract nested keys from inside the columns dictionary block
    #         cols_dict = table_info.get("columns", {})
    #         if isinstance(cols_dict, dict):
    #             for col_name in cols_dict.keys():
    #                 keywords.add(str(col_name).lower())
    #                 # Split column names containing underscores to capture separate words
    #                 if "_" in col_name:
    #                     keywords.update(col_name.lower().split("_"))
                        
    #     keywords = {kw.strip() for kw in keywords if len(kw.strip()) > 1}
    #     return keywords

    

    # def fast_sql_check(self, user_query):
    #     clean_query = re.sub(r'[^\w\s]', ' ', user_query.lower())
    #     query_words = set(clean_query.split())
    #     matched_keywords = query_words.intersection(self.schema_keywords)
    #     return list(matched_keywords)

    # def has_sql_intent(self, user_query):
    #     query_lower = user_query.lower()
    #     return any(pattern in query_lower for pattern in self.sql_patterns)   

    # --- Placeholder Functions (Kept exactly as requested) ---
    def generate_sql(self, natural_language_query, database_schema):
        pass
    
    def chat_completion(self, messages, model_config):
        pass
    
    def generate_insights(self, data, context):
        pass
    
    def validate_sql(self, sql_query, database_schema):
        pass

    def _table_to_markdown(self, rows):
        """Turns a list-of-rows table into a clean markdown table string."""
        if not rows:
            return ""
        header = rows[0]
        lines = ["| " + " | ".join(str(c or "") for c in header) + " |"]
        lines.append("|" + "|".join(["---"] * len(header)) + "|")
        for row in rows[1:]:
            lines.append("| " + " | ".join(str(c or "") for c in row) + " |")
        return "\n".join(lines)

    def _extract_pdf_tables(self, file_path, document_code):
        import pdfplumber
        from langchain_core.documents import Document
        import uuid

        table_cfg = self.rag_config["chunking"]["table"]
        chunks = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    for table_index, table in enumerate(page.extract_tables()):
                        if not table or len(table) < 2:
                            continue
                        table_id = f"{document_code}_table_p{page_num}_{table_index}_{uuid.uuid4().hex[:6]}"
                        header, body_rows = table[0], table[1:]
                        if len(body_rows) <= table_cfg["keep_whole_if_under_rows"]:
                            chunks.append(Document(
                                page_content=self._table_to_markdown(table),
                                metadata={"document_code": document_code, "table_id": table_id,
                                          "source": "pdf_table", "page": page_num}
                            ))
                        else:
                            group_size = table_cfg["keep_whole_if_under_rows"]
                            for i in range(0, len(body_rows), group_size):
                                group = body_rows[i:i + group_size]
                                piece_rows = ([header] + group) if table_cfg["repeat_headers_on_split"] else group
                                chunks.append(Document(
                                    page_content=self._table_to_markdown(piece_rows),
                                    metadata={"document_code": document_code, "table_id": table_id,
                                              "source": "pdf_table", "page": page_num}
                                ))
        except Exception as e:
            print(f"⚠️ [RAG] Could not extract PDF tables from {file_path}: {e}")
        return chunks

    def _extract_docx_tables(self, file_path, document_code):
        from docx import Document as DocxDocument
        from langchain_core.documents import Document
        import uuid

        table_cfg = self.rag_config["chunking"]["table"]
        chunks = []
        try:
            doc = DocxDocument(file_path)
            for table_index, table in enumerate(doc.tables):
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                if len(rows) < 2:
                    continue
                table_id = f"{document_code}_table_{table_index}_{uuid.uuid4().hex[:6]}"
                header, body_rows = rows[0], rows[1:]
                if len(body_rows) <= table_cfg["keep_whole_if_under_rows"]:
                    chunks.append(Document(
                        page_content=self._table_to_markdown(rows),
                        metadata={"document_code": document_code, "table_id": table_id, "source": "docx_table"}
                    ))
                else:
                    group_size = table_cfg["keep_whole_if_under_rows"]
                    for i in range(0, len(body_rows), group_size):
                        group = body_rows[i:i + group_size]
                        piece_rows = ([header] + group) if table_cfg["repeat_headers_on_split"] else group
                        chunks.append(Document(
                            page_content=self._table_to_markdown(piece_rows),
                            metadata={"document_code": document_code, "table_id": table_id, "source": "docx_table"}
                        ))
        except Exception as e:
            print(f"⚠️ [RAG] Could not extract DOCX tables from {file_path}: {e}")
        return chunks

    def _extract_pdf_images(self, file_path, document_code):
        import fitz
        import base64
        from langchain_core.documents import Document
        import uuid

        image_cfg = self.rag_config["ingestion"]["images"]
        chunks = []
        try:
            pdf_doc = fitz.open(file_path)
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    if len(image_bytes) < image_cfg["min_image_size_bytes"]:
                        continue
                    b64_image = base64.b64encode(image_bytes).decode("utf-8")
                    caption = self._caption_image(b64_image, base_image["ext"])
                    if caption:
                        image_id = f"{document_code}_image_p{page_num}_{img_index}_{uuid.uuid4().hex[:6]}"
                        chunks.append(Document(
                            page_content=caption,
                            metadata={"document_code": document_code, "image_id": image_id,
                                      "source": "pdf_image", "page": page_num}
                        ))
            pdf_doc.close()
        except Exception as e:
            print(f"⚠️ [RAG] Could not extract PDF images from {file_path}: {e}")
        return chunks

    def _caption_image(self, b64_image, image_ext):
        try:
            vision_llm = ChatOpenAI(model="gpt-4o-mini", temperature=self.rag_config["generation"]["temperature"], openai_api_key=os.getenv("OPENAI_API_KEY"))
            message = HumanMessage(content=[
                {"type": "text", "text": "Describe this image in 2-3 sentences, focusing on any data, numbers, chart values, or important details a business user would need to know."},
                {"type": "image_url", "image_url": {"url": f"data:image/{image_ext};base64,{b64_image}"}}
            ])
            response = vision_llm.invoke([message])
            return response.content.strip()
        except Exception as e:
            print(f"⚠️ [RAG] Could not caption image: {e}")
            return None

    def _generate_hyde_document(self, query):
        """Writes a short hypothetical answer to search with, instead of the raw question."""
        try:
            hyde_llm = ChatOpenAI(
                model=self.rag_config["retrieval"]["hyde"]["model"],
                temperature=0.3,
                openai_api_key=os.getenv("OPENAI_API_KEY")
            )
            prompt = (
                "Write a short, plausible-sounding paragraph that could be the answer to "
                "this question. It does not need to be factually correct - it only needs "
                "to read like a real document that would contain the answer. Do not "
                "include any disclaimers, just write the paragraph.\n\n"
                f"Question: {query}"
            )
            response = hyde_llm.invoke([HumanMessage(content=prompt)])
            return response.content.strip()
        except Exception as e:
            print(f"⚠️ [RAG] HyDE generation failed, using raw query instead: {e}")
            return query

    def _generate_query_variations(self, query):
        """Asks the LLM to reword the question a few different ways, for multi-query retrieval."""
        mq_cfg = self.rag_config["retrieval"]["multi_query"]
        try:
            mq_llm = ChatOpenAI(
                model=mq_cfg["model"],
                temperature=0.5,
                openai_api_key=os.getenv("OPENAI_API_KEY")
            )
            prompt = (
                f"Reword this question {mq_cfg['num_variations']} different ways, keeping "
                "the same meaning. Return only the numbered list, nothing else.\n\n"
                f"Question: {query}"
            )
            response = mq_llm.invoke([HumanMessage(content=prompt)])
            lines = [l.strip() for l in response.content.strip().split("\n") if l.strip()]
            cleaned = []
            for line in lines:
                line = re.sub(r'^[\d\.\)\-\s]+', '', line).strip()
                if line:
                    cleaned.append(line)
            return cleaned[:mq_cfg["num_variations"]]
        except Exception as e:
            print(f"⚠️ [RAG] Multi-query generation failed, using original query only: {e}")
            return []

    def _merge_and_dedupe_docs(self, doc_lists):
        """Combines results from multiple searches, dropping exact duplicate chunk content."""
        seen = set()
        merged = []
        for docs in doc_lists:
            for doc in docs:
                key = doc.page_content.strip()
                if key not in seen:
                    seen.add(key)
                    merged.append(doc)
        return merged

    # --- Core Processing Logic ---

    def process_to_embeddings(self, file_path, document_code=None):
        """
        Loads document, creates chunks, generates embeddings, and stores in Qdrant.
        """
        if not os.path.exists(file_path):
            return {"error": f"File not found at {file_path}"}

        ext = os.path.splitext(file_path)[-1].lower()
        
        try:
            # 1. Loader Selection
            if ext == ".pdf":
                loader = PyPDFLoader(file_path)
            elif ext in [".docx", ".doc"]:
                loader = Docx2txtLoader(file_path)
            elif ext == ".md":
                loader = UnstructuredMarkdownLoader(file_path)
            elif ext == ".rtf":
                loader = UnstructuredRTFLoader(file_path)
            else:
                loader = TextLoader(file_path)

            # Load the data
            documents = loader.load()

            # Extract tables and images separately, gated by rag_config.yaml so this
            # is a no-op until those settings are turned on.
            # NOTE: the normal text loader above will still also pull table content
            # as garbled plain text - some duplication is expected, not a bug.
            table_chunks = []
            if self.rag_config["chunking"]["table"]["enabled"]:
                if ext == ".pdf":
                    table_chunks = self._extract_pdf_tables(file_path, document_code)
                elif ext in [".docx", ".doc"]:
                    table_chunks = self._extract_docx_tables(file_path, document_code)

            image_chunks = []
            if ext == ".pdf" and self.rag_config["ingestion"]["images"]["extract_from_pdf"]:
                image_chunks = self._extract_pdf_images(file_path, document_code)

            # 2. Chunking
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.rag_config["chunking"]["chunk_size"],
                chunk_overlap=self.rag_config["chunking"]["chunk_overlap"]
            )
            chunks = text_splitter.split_documents(documents)

            # Add metadata to identify chunks by document code
            for chunk in chunks:
                chunk.metadata["document_code"] = document_code

            chunks.extend(table_chunks)
            chunks.extend(image_chunks)

            # 3. Store in Qdrant Vector Database
            # This handles both embedding and storage in one go
            QdrantVectorStore.from_documents(
                documents=chunks,
                embedding=self.embeddings_model,
                url=self.qdrant_url,
                collection_name=self.collection_name,
                force_recreate=False
            )
            
            # Return the exact success message for your frontend
            return {
                "status": "success",
                "message": "Embeddings are stored in vector database successfully",
                "chunk_count": len(chunks)
            }

        except Exception as e:
            return {"error": f"Processing failed: {str(e)}"}

    
    # def perform_intent_analysis(self, user_query, model_name, custom_key):
    #     """Helper to handle Step 2 intent analysis dynamically."""
    #     analysis_prompt = f"Describe the analysis of this query in one short sentence: '{user_query}'. Output only the sentence."
        
    #     try:
    #         # Check if it's one of your pre-configured GPT models
    #         if model_name in self.models:
    #             messages = [HumanMessage(content=analysis_prompt)]
    #             return self.models[model_name].invoke(messages).content.strip()
            
    #         # Logic for other providers (Claude, Gemini, etc.) would go here...
    #         # ... (same logic as the snippet I provided earlier)
            
    #     except Exception:
    #         return "Analyzing natural language inquiry for document matching modules."
        
    # def perform_intent_analysis(self, user_query, model_name, custom_key):
    #     """Helper to handle Step 2 intent analysis dynamically."""
    #     analysis_prompt = f"Describe the analysis of this query in one short sentence: '{user_query}'. Output only the sentence."
    #     messages = [HumanMessage(content=analysis_prompt)]
        
    #     try:
    #         # 1. Check if the model_name is one of your pre-configured GPT models
    #         if model_name in self.models:
    #             # If a custom_key is provided, we override the default key for this specific call
    #             if custom_key:
    #                 # Temporary instance with the custom key
    #                 temp_llm = ChatOpenAI(model=model_name, temperature=0, openai_api_key=custom_key)
    #                 return temp_llm.invoke(messages).content.strip()
                
    #             # Otherwise, use the pre-initialized model
    #             return self.models[model_name].invoke(messages).content.strip()
            
    #         # 2. If it's not a pre-configured model, it returns the default
    #         return "Analyzing natural language inquiry for document matching modules."
            
    #     except Exception as e:
    #         print(f"Analysis Error: {e}")
    #         return "Analyzing natural language inquiry for document matching modules."    

    def answer_from_docs(self, user_query, model_name,session_id=1,custom_key='',system_instructions=''):
        """
        Retrieves relevant chunks from Qdrant and updates the live steps 
        using the new structured event payload layout.
        """
        rag_chain_of_thought = []
        session_id = str(session_id)
        
        # ========================================================
       
        # HELPER FUNCTION TO STREAM DICTIONARY OBJECTS TO THE NEW UI
        # ========================================================
        def push_rag_event(event_type, title, description):
            ui_title = title
            if title == "Query Intent Analysis":
                ui_title = "Context Intent Analysis"
            
            payload = {
                "event": event_type,
                "title": title,
                "description": description,
                "is_sql": False
            }

            if event_type == "start":
                rag_chain_of_thought.append(f"{title} - {description}")
            
            # Pushes the required map data structure down to StepStreamManager
            stream_manager.push_step(session_id, payload, is_sql=False)
            time.sleep(0.3) # Give frontend rendering engine time to execute animations
        
        # ========================================================

        try:
            # ========================================================
            # STEP 1: INQUIRY RECEIVED
            # ========================================================
            push_rag_event("start", "Query Received", f"'{user_query}'")
            push_rag_event("complete", "Query Received", f"'{user_query}'")
           
            # Old implementation: stream_manager.push_step(session_id, "Step 1: Document Query Parsing...", False)
            push_rag_event("start", "Query Parsing", f"Analyzing raw unstructured prompt: '{user_query}'")
            # --- CODE CHANGE END ---
            
            client = QdrantClient(url=self.qdrant_url)
            vector_store = QdrantVectorStore(
                client=client,
                collection_name=self.collection_name,
                embedding=self.embeddings_model
            )
            
            # --- CODE CHANGE START ---
            push_rag_event("complete", "Query Parsing", "Query successfully processed and converted to dense vectors.")
            # --- CODE CHANGE END ---

            # ========================================================
            # STEP 2: INTENT ANALYSIS
            # ========================================================

            # push_rag_event("start", "Context Intent Analysis", "Analyzing...")
            # analysis_text = self.perform_intent_analysis(user_query, model_name, custom_key)
            # push_rag_event("complete", "Context Intent Analysis", analysis_text)
            # --- CODE CHANGE START ---
            # ========================================================
            # STEP 2: INTENT ANALYSIS
            # ========================================================
            push_rag_event("start", "Query Intent Analysis", f"Querying {model_name} instance for contextual definition profiles...")
            
            analysis_prompt = f"Describe the analysis of this query in one short sentence: '{user_query}'. Output only the sentence."
            
            try:
                # 1. CLOUD OPENAI INSTANCES FROM THE DICTIONARY
                if model_name in self.models:
                    messages = [HumanMessage(content=analysis_prompt)]
                    ai_response = self.models[model_name].invoke(messages)
                    analysis_text = ai_response.content.strip()
                
                # 2. DYNAMIC CUSTOM API ROUTING (Claude, Gemini, DeepSeek, etc.)
                elif str(model_name).startswith("api://"):
                    actual_model = model_name.replace("api://", "").lower()
                    messages = [HumanMessage(content=analysis_prompt)]
                    
                    if "claude" in actual_model:
                        from langchain_anthropic import ChatAnthropic
                        dynamic_llm = ChatAnthropic(
                            model=actual_model,
                            temperature=self.rag_config["generation"]["temperature"],
                            anthropic_api_key=custom_key if custom_key else os.getenv("ANTHROPIC_API_KEY")
                        )
                    elif "gemini" in actual_model:
                        from langchain_google_genai import ChatGoogleGenerativeAI
                        dynamic_llm = ChatGoogleGenerativeAI(
                            model=actual_model,
                            temperature=self.rag_config["generation"]["temperature"],
                            google_api_key=custom_key if custom_key else os.getenv("GOOGLE_API_KEY")
                        )
                    elif "deepseek" in actual_model:
                        dynamic_llm = ChatOpenAI(
                            model=actual_model,
                            temperature=self.rag_config["generation"]["temperature"],
                            openai_api_key=custom_key if custom_key else os.getenv("DEEPSEEK_API_KEY"),
                            openai_api_base="https://api.deepseek.com/v1"
                        )
                    else:  # Custom GPT models
                        dynamic_llm = ChatOpenAI(
                            model=actual_model,
                            temperature=self.rag_config["generation"]["temperature"],
                            openai_api_key=custom_key if custom_key else self.openai_key
                        )
                    
                    ai_response = dynamic_llm.invoke(messages)
                    analysis_text = ai_response.content.strip()

                # 3. DYNAMIC OLLAMA ROUTING
                elif str(model_name).startswith("ollama://") or model_name == "llama3":
                    actual_model = model_name.replace("ollama://", "") if str(model_name).startswith("ollama://") else "llama3"
                    cot_payload = {
                        "model": actual_model,
                        "prompt": analysis_prompt,
                        "stream": False,
                        "options": {"temperature": self.ollama_config["temperature"]}
                    }
                    cot_res = requests.post(self.ollama_config["url"], json=cot_payload, timeout=60)
                    analysis_text = cot_res.json().get("response", "").strip()
                
                else:
                    analysis_text = "Analyzing natural language inquiry for document matching modules."

            except Exception as e:
                print(f"⚠️ Step 2 Dynamic Analysis Fallback Trace: {e}")
                analysis_text = "Analyzing natural language inquiry for document matching modules."
            
            push_rag_event("complete", "Query Intent Analysis", analysis_text)

            #################################################
            # push_rag_event("start", "Context Intent Analysis", "Querying local LLM instance for contextual definition profiles...")
            # # --- CODE CHANGE END ---
            # try:
            #      cot_payload = {
            #         "model": self.ollama_config["model"],
            #         "prompt": f"Describe the analysis of this query in one short sentence: '{user_query}'. Output only the sentence.",
            #         "stream": False,
            #          "options": {"temperature": self.ollama_config["temperature"]}
            #     }
            #      cot_res = requests.post(self.ollama_config["url"], json=cot_payload, timeout=300)
            #      analysis_text = cot_res.json().get("response", "Analyzing document index for relevant parameters.").strip()
            # except Exception:
            #      analysis_text = "Analyzing natural language inquiry for document matching modules."
            
            # # # --- CODE CHANGE START ---
            # push_rag_event("complete", "Context Intent Analysis", analysis_text)
            # # --- CODE CHANGE END ---

            # ========================================================
            # STEP 3: KNOWLEDGE BASE RETRIEVAL
            # ========================================================
            # --- CODE CHANGE START ---
            push_rag_event("start", "Vector Space Search", "Scanning local Qdrant collections for high-probability matching fragments...")
            # --- CODE CHANGE END ---
            
            retrieval_cfg = self.rag_config["retrieval"]
            search_queries = [user_query]

            if retrieval_cfg["multi_query"]["enabled"]:
                search_queries.extend(self._generate_query_variations(user_query))

            if retrieval_cfg["hyde"]["enabled"]:
                search_queries = [self._generate_hyde_document(q) for q in search_queries]

            if len(search_queries) > 1:
                doc_lists = [vector_store.similarity_search(q, k=retrieval_cfg["top_k"]) for q in search_queries]
                docs = self._merge_and_dedupe_docs(doc_lists)
            else:
                docs = vector_store.similarity_search(search_queries[0], k=retrieval_cfg["top_k"])
            context_text = "\n\n".join([doc.page_content for doc in docs])
            
            retrieval_msg = f"Successfully matched and isolated {len(docs)} high-relevance documentation segments."
            
            # --- CODE CHANGE START ---
            push_rag_event("complete", "Vector Space Search", retrieval_msg)
            # --- CODE CHANGE END ---

            if not context_text:
                stream_manager.push_step(session_id, "DONE", is_sql=False)
                return {
                    "answer": "I couldn't find any relevant information in the uploaded documents.",
                    "sql": None,
                    "table": [],
                    "chart": {},
                    "rag_chain_of_thought": rag_chain_of_thought
                }

            # ========================================================
            # STEP 4: RESPONSE SYNTHESIS
            # ========================================================
            # --- CODE CHANGE START ---
            push_rag_event("start", "Output Synthesis", f"Processing matrices through {model_name} to compile answers...")
            # --- CODE CHANGE END ---

            system_prompt = (
                "You are Saarthi AI, a helpful assistant. Answer the user's question "
                "using ONLY the following context. If the answer is not in the context, "
                "politely say you don't know based on the documents.\n\n"
                f"CONTEXT:\n{context_text}"
            )
            if system_instructions.strip():
                system_prompt += f"\n\n[CRITICAL PERSONA AND CUSTOM FORMATTING RULES]:\n{system_instructions}"

            if model_name == "llama3":
                print("🦙 Routing payload to local Ollama [llama3] container layer...")
                ollama_prompt = f"{system_prompt}\n\nUSER QUESTION:\n{user_query}"
                
                response = requests.post(
                    self.ollama_config["url"],
                    json={
                        "model": "llama3",  # Forces local container system instance call
                        "prompt": ollama_prompt,
                        "stream": False,
                        "keep_alive": "30m",
                        "options": {
                            "temperature": self.ollama_config["temperature"],
                            "num_ctx": self.ollama_config["num_ctx"],
                            "num_thread": 4
                        }
                    },
                    timeout=self.ollama_config["timeout"]
                )
                response.raise_for_status()
                final_answer = response.json().get("response", "").strip()
                
            elif model_name == "gpt-4o":
                print("🔥 Routing payload to cloud production instance: GPT-4o Premium...")
                #self.llm.model_name = "gpt-4o"
                #llm_instance = ChatOpenAI(
                #    model="gpt-4o",
                #    temperature=0,
                #    openai_api_key=os.getenv("OPENAI_API_KEY")
                #)
                messages = [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_query)
                ]
                #ai_response = self.llm.invoke(messages)
                ai_response = self.models["gpt-4o"].invoke(messages)
                final_answer = ai_response.content
                
            elif model_name == "gpt-4o-mini":
                print("🤖 Routing payload to cloud production instance: GPT-4o Mini...")
                #self.llm.model_name = "gpt-4o-mini"
                #llm_instance = ChatOpenAI(
                #    model="gpt-4o-mini",
                #    temperature=0,
                #    openai_api_key=os.getenv("OPENAI_API_KEY")
                #)
                messages = [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_query)
                ]
                #ai_response = self.llm.invoke(messages)
                ai_response = self.models["gpt-4o-mini"].invoke(messages)
                final_answer = ai_response.content
            
            elif str(model_name).startswith("api://"):
                actual_model = model_name.replace("api://", "").lower()
                print(f"🌐 Dynamic RAG Routing payload to Custom Cloud API model: {actual_model}")
                
                messages = [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_query)
                ]

                # 1. ANTHROPIC CLAUDE MODELS
                if "claude" in actual_model:
                    from langchain_anthropic import ChatAnthropic
                    dynamic_llm = ChatAnthropic(
                        model=actual_model,
                        temperature=self.rag_config["generation"]["temperature"],
                        anthropic_api_key=custom_key if custom_key else os.getenv("ANTHROPIC_API_KEY")
                    )
                    ai_response = dynamic_llm.invoke(messages)
                    final_answer = ai_response.content

                # 2. GOOGLE GEMINI MODELS
                elif "gemini" in actual_model:
                    from langchain_google_genai import ChatGoogleGenerativeAI
                    dynamic_llm = ChatGoogleGenerativeAI(
                        model=actual_model,
                        temperature=self.rag_config["generation"]["temperature"],
                        google_api_key=custom_key if custom_key else os.getenv("GOOGLE_API_KEY")
                    )
                    ai_response = dynamic_llm.invoke(messages)
                    final_answer = ai_response.content

                # 3. DEEPSEEK MODELS
                elif "deepseek" in actual_model:
                    dynamic_llm = ChatOpenAI(
                        model=actual_model,
                        temperature=self.rag_config["generation"]["temperature"],
                        openai_api_key=custom_key if custom_key else os.getenv("DEEPSEEK_API_KEY"),
                        openai_api_base="https://api.deepseek.com/v1"
                    )
                    ai_response = dynamic_llm.invoke(messages)
                    final_answer = ai_response.content

                # 4. EXPLICIT CUSTOM OPENAI MODELS
                elif "gpt" in actual_model or "openai" in actual_model:
                    dynamic_llm = ChatOpenAI(
                        model=actual_model,
                        temperature=self.rag_config["generation"]["temperature"],
                        openai_api_key=custom_key if custom_key else self.openai_key
                    )
                    ai_response = dynamic_llm.invoke(messages)
                    final_answer = ai_response.content
                else:
                    raise ValueError(
                        f"Custom cloud provider mapping failed: Identifier '{actual_model}' "
                        f"does not match any recognized provider keyword (claude, gemini, deepseek, gpt)."
                    )

            elif str(model_name).startswith("ollama://"):
                actual_model = model_name.replace("ollama://", "")
                print(f"📦 Dynamic RAG Routing payload to Custom Local Ollama model: {actual_model}")
                ollama_prompt = f"{system_prompt}\n\nUSER QUESTION:\n{user_query}"
                
                response = requests.post(
                    self.ollama_config["url"],
                    json={
                        "model": actual_model,  # Directly maps raw string to local runtime container target
                        "prompt": ollama_prompt,
                        "stream": False,
                        "keep_alive": "30m",
                        "options": {
                            "temperature": self.ollama_config["temperature"],
                            "num_ctx": self.ollama_config["num_ctx"],
                            "num_thread": 4
                        }
                    },
                    timeout=self.ollama_config["timeout"]
                )
                response.raise_for_status()
                final_answer = response.json().get("response", "").strip()    
                
            else:
                raise ValueError(f"Requested model '{model_name}' has no active route handler configuration.")
            # --- CODE CHANGE START ---
            push_rag_event("complete", "Output Synthesis", f"generated response through {model_name}")
            # --- CODE CHANGE END ---

            # Final execution loop boundary closeout
            stream_manager.push_step(session_id, "DONE", is_sql=False)

            return {
                "answer": final_answer,
                "sql": None,  
                "table": [],
                "chart": {},
                "rag_chain_of_thought": rag_chain_of_thought
            }

        except Exception as e:
            print(f"Retrieval Error Trace: {str(e)}")
            stream_manager.push_step(session_id, "DONE", is_sql=False)
            return {
                "answer": "The system encountered an error while searching the document database.",
                "sql": None,
                "table": [],
                "chart": {},
                "rag_chain_of_thought": rag_chain_of_thought
            } 

_shared_llm_service = LLMService()

def answer_from_docs(user_query, model_name, session_id=1, custom_key='', system_instructions=''):
    """
    Top-level module function mapping so your orchestrator router can import 
    it cleanly without needing class-level structural instantiation overhead.
    """
    return _shared_llm_service.answer_from_docs(
        user_query=user_query,
        model_name=model_name,
        session_id=session_id,
        custom_key=custom_key,
        system_instructions=system_instructions
    )   
       

#     def get_smart_response(self, user_query,model_name, session_id=1,custom_key=''):
#         """
#         Layered high-speed query execution routing pipeline.
#         Uses phi3:mini directly to categorize user query path intent.
#         """
#         try:
#             print("\n" + "=" * 60)
#             print(f"🧠 SMART ROUTER PROCESSING QUERY: {user_query}")

#             # ----------------------------------------------------
#             # LAYER 1: FAST HEURISTIC PATTERN PASS
#             # ----------------------------------------------------
#             matched_keywords = self.fast_sql_check(user_query)
#             sql_intent = self.has_sql_intent(user_query)

#             lower_query = user_query.lower().strip()
            
#             # 1. Broad conversational/informational phrase triggers
#             rag_phrases = ["what is", "what does", "tell me about", "explain", "meaning of", "who is", "how do i", "where can i"]
            
#             # 2. Key operational indicators that clearly target structured computational data
#             sql_indicators = ["total", "how many", "count", "sum", "average", "avg", "list", "top", "maximum", "max", "highest", "lowest"]
            
#             has_rag_phrase = any(phrase in lower_query for phrase in rag_phrases)
#             has_sql_indicator = any(indicator in lower_query for indicator in sql_indicators)
            
#             # If a query matches a conversational pattern but contains NO data computation keywords 
#             # and triggers no specific database table keywords, intercept it immediately as a RAG document request.
#             if has_rag_phrase and not has_sql_indicator:
#             #if has_rag_phrase and not (matched_keywords or has_sql_indicator):
#                 print("📄 [INTELLIGENT RAG GUARD] -> Pure descriptive intent detected. Routing directly to RAG.")
#                 rag_res = self.answer_from_docs(user_query, model_name=model_name,session_id=session_id,custom_key=custom_key)
#                 steps_list = rag_res.get("rag_chain_of_thought", [])
#                 return {
#                     "answer": rag_res.get("answer"),
#                     "sql": None,
#                     "table": [],
#                     "chart": {},
#                     "steps": steps_list
#                 }

#             # ----------------------------------------------------
#             # LAYER 2: FAST HEURISTIC PATTERN PASS FOR PURE SQL
#             # ----------------------------------------------------
#             if matched_keywords and sql_intent:
#                 print("🚀 [FAST PATH TRIGGERED] -> Confirmed structured operational query. Routing straight to SQL.")
#                 full_result = run_data_bridge_agent(user_query, session_id=session_id,model_name=model_name,custom_key=custom_key)
#                 return full_result["chat_ui"]



#             router_prompt = f"""
# You are an intelligent query router for an enterprise AI system.

# DATABASE SCHEMA:
# {self.schema_str}

# USER QUERY:
# "{user_query}"

# YOUR JOB:
# Decide if this query should be answered from:
# 1. SQL - database tables shown above
# 2. RAG - uploaded documents/PDFs/text files

# RULES:
# - If the query is asking for data that EXISTS as a column in the schema above, choose SQL.
# - If the query is asking for information NOT found in the schema (explanations, descriptions, policies, document content,summaries), choose RAG.
# - If the same word (like "company") exists both in schema AND could be in documents, look at the INTENT:
#   * "List all companies" → SQL (wants rows/records)
#   * "What does the company do?" → RAG (wants description)
#   * "How many companies?" → SQL (wants count)
#   * "Tell me about the company" → RAG (wants explanation)

# Respond with ONLY: SQL or RAG.
# """

#             response = requests.post(
#                 self.ollama_config["url"],
#                 json={
#                     "model": self.ollama_config["model"],  # ⚡ Explicitly runs phi3:mini
#                     "prompt": router_prompt,
#                     "stream": False,
#                     "keep_alive": "30m",
#                     "options": {
#                         "temperature": 0.0,
#                         "num_ctx": 8192,
#                         "num_predict": 3,
#                         "num_thread": 4
#                     }
#                 },
#                 timeout=60
#             )
#             response.raise_for_status()
#             res_text = response.json().get("response", "").strip().upper()
#             print(f"🧠 Router LLM Decision: {res_text}")

#             # ----------------------------------------------------
#             # LAYER 3: PATH SPECIFIC EXECUTION ASSIGNMENT
#             # ----------------------------------------------------
#             if "SQL" in res_text:
#                 print("📊 ROUTING DIRECTION -> DATA BRIDGE LANGRAPH NETWORK")

#                 import time
#                 time.sleep(0.2)
#                 stream_manager.push_step(
#                     session_id, 
#                     "Router Analysis - Fast Path Heuristic Triggered: Bypassing router LLM.", 
#                     is_sql=True)
#                 full_result = run_data_bridge_agent(user_query, session_id=session_id,model_name=model_name)
#                 stream_manager.push_step(session_id, "DONE", is_sql=True)

#                 if isinstance(full_result.get("chat_ui"), dict) and "Error at error_diagnosis" in full_result["chat_ui"].get("answer", ""):
#                     full_result["chat_ui"]["answer"] = (
#                         "I encountered an issue while generating the database query. Please check if the requested column or table exists in your schema."
#                     )

#                 return full_result["chat_ui"]
#             else:
#                 print("📄 ROUTING DIRECTION -> UNSTRUCTURED KNOWLEDGE DATABASE RAG")
#                 #return self.answer_from_docs(user_query, session_id=session_id)
#                 rag_res = self.answer_from_docs(user_query,model_name=model_name,session_id=session_id,custom_key=custom_key)
#                 steps_list = rag_res.get("rag_chain_of_thought", [])
#                 #for step in steps_list:
#                 #    stream_manager.push_step(session_id, step, is_sql=False)

#                 return {
#                     "answer": rag_res.get("answer"),
#                     "sql": None,
#                     "table": [],
#                     "chart": {},
#                     "steps": steps_list  # Maps rag_chain_of_thought to steps
#                 }
#         except requests.exceptions.Timeout:
#             print("⏰ Router processing limit reached -> Fallback to RAG")
#             rag_res = self.answer_from_docs(user_query, model_name=model_name,session_id=session_id,custom_key=custom_key)
#             steps_list = rag_res.get("rag_chain_of_thought", [])
            
#             for step in steps_list:
#                 stream_manager.push_step(session_id, step, is_sql=False)
#             stream_manager.push_step(session_id, "DONE", is_sql=False)    
#             return {
#                 "answer": rag_res.get("answer"),
#                 "sql": None,
#                 "table": [],
#                 "chart": {},
#                 "steps": steps_list
#             }
        
#         except Exception as e:
#             import traceback
#             print("\n❌ [CRITICAL SQL PIPELINE FAILURE] The database agent crashed!")
#             traceback.print_exc()
#             stream_manager.push_step(
#                 session_id, 
#                 "❌ Pipeline Execution Failure detected in Router Layer.", 
#                 is_sql=True
#             )
#             stream_manager.push_step(
#                 session_id, 
#                 f"⚠️ Diagnosed Error: {str(e)}", 
#                 is_sql=True
#             )
#             stream_manager.push_step(session_id, "DONE", is_sql=True)
            
#             return {
#                 "answer": f"SQL Execution Error: {str(e)}",
#                 "sql": None,
#                 "table": [],
#                 "chart": {},
#                 "steps": [f"Pipeline failed at router layer: {str(e)}"]
#             }

        
                  