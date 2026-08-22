#!/usr/bin/env python3
"""
generate_guideline_docs.py
===========================
Generates the "unstructured document" side of the lending demo: a small
set of internal NBFC policy documents (.docx) written in-house as
practitioner-style summaries of well-known, publicly discussed Indian NBFC
regulatory themes (RBI Master Directions on NBFCs, Fair Practices Code,
KYC/AML norms, IRAC/NPA classification, digital lending, grievance
redressal). These are original summaries for demo/testing purposes, NOT a
reproduction of any official RBI publication, and should not be treated as
a compliance reference - only as sample content for exercising Saarthi's
document-RAG pipeline against realistic-looking policy text.

Usage:
    python generate_guideline_docs.py
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output", "guidelines")

NAVY = RGBColor(0x1F, 0x38, 0x64)


def add_title(doc: Document, title: str, doc_code: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Internal Policy Reference | {doc_code} | LendWell NBFC Ltd. (synthetic demo entity)")
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "This is a synthetic, practitioner-style summary prepared for software demo/testing purposes. "
        "It is not an official RBI publication and must not be used as a compliance reference."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0xA0, 0x30, 0x30)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str):
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph()


# ---------------------------------------------------------------------
# Document 1: Fair Practices Code
# ---------------------------------------------------------------------

def build_fair_practices_code() -> Document:
    doc = Document()
    add_title(doc, "Fair Practices Code for Lending Operations", "FPC-001 v3.2")

    add_heading(doc, "1. Purpose and Applicability")
    add_body(doc,
        "This Code sets out the minimum standards of fair, transparent and non-coercive conduct that "
        "LendWell NBFC Ltd. follows across the loan lifecycle - from advertising and origination through "
        "servicing, collections and closure - in line with the spirit of RBI's guidance for NBFCs on fair "
        "lending conduct. It applies to all branches, digital channels, and third-party agents/DSAs "
        "engaged by the company.")

    add_heading(doc, "2. Loan Application and Processing")
    add_bullets(doc, [
        "All loan application forms disclose the information required for an informed decision, including "
        "fees, interest rate basis (fixed/floating), penal charges and prepayment terms.",
        "An acknowledgement is issued for every application received, indicating the expected turnaround "
        "time for a decision.",
        "Rejections are communicated in writing along with the specific reason(s) for rejection.",
        "No discrimination is made on grounds of gender, marital status, religion, caste or disability in "
        "the loan appraisal process; only creditworthiness-related factors are used.",
    ])

    add_heading(doc, "3. Loan Appraisal and Sanction")
    add_body(doc,
        "The sanction letter/loan agreement sets out the amount sanctioned, annualised interest rate, "
        "method of application (reducing balance), tenure, EMI schedule, processing fee, and all other "
        "charges. A copy of the loan agreement and enclosures is furnished to every borrower at the time "
        "of disbursement.")

    add_heading(doc, "4. Disbursement, Changes in Terms and Post-Disbursement Conduct")
    add_bullets(doc, [
        "Any change in interest rate, charges or terms is notified to the borrower in advance, in a "
        "language understood by the borrower, through the agreed communication channel.",
        "Decisions to recall/accelerate a loan or seek additional security are taken in consonance with the "
        "loan agreement and only after due notice.",
        "On repayment of all dues, all security documents are released within 30 days, subject to any "
        "lawful right of set-off the company may have.",
    ])

    add_heading(doc, "5. Interest Rate Policy")
    add_body(doc,
        "The Board-approved interest rate model considers cost of funds, margin, risk premium for the "
        "borrower/product category, and is reviewed at least annually. The range of interest rates charged "
        "for a given product/category, and the rationale for the differential, is disclosed on the "
        "company's website and in loan documentation. Penal charges for delayed payment are levied as a "
        "reasonable, disclosed percentage and are not compounded as additional interest.")

    add_heading(doc, "6. Collection of Dues and Recovery Practices")
    add_bullets(doc, [
        "Collections staff and empanelled recovery agents identify themselves and produce authorisation "
        "on request.",
        "Contact with borrowers for recovery purposes is restricted to 8:00 AM to 7:00 PM unless the "
        "borrower has agreed otherwise.",
        "Recovery agents do not use threatening, intimidating, or abusive language, and do not contact "
        "the borrower's referees or family members except to trace an untraceable borrower.",
        "All borrower grievances relating to recovery conduct are logged and reviewed under the Grievance "
        "Redressal Policy (see GRP-004).",
    ])

    add_heading(doc, "7. Responsibility of the Board")
    add_body(doc,
        "The Board of Directors reviews compliance with this Code and the functioning of the grievance "
        "redressal mechanism at least once every financial year, and lays down the appropriate grievance "
        "redressal mechanism within the organisation.")

    return doc


# ---------------------------------------------------------------------
# Document 2: KYC / AML Policy
# ---------------------------------------------------------------------

def build_kyc_aml_policy() -> Document:
    doc = Document()
    add_title(doc, "Know Your Customer (KYC) and Anti-Money Laundering (AML) Policy", "KYC-002 v4.0")

    add_heading(doc, "1. Objective")
    add_body(doc,
        "This policy establishes the customer identification, due diligence and monitoring standards "
        "LendWell follows so that the company is not used, intentionally or unintentionally, for money "
        "laundering or terror financing, and to comply with the spirit of RBI's Master Direction on KYC "
        "and the Prevention of Money Laundering Act (PMLA), 2002 and its Rules.")

    add_heading(doc, "2. Customer Acceptance and Risk Categorisation")
    add_bullets(doc, [
        "Every customer is categorised as Low, Medium or High risk at onboarding based on occupation, "
        "income profile, geographical location, mode of payment, and nature of business (for self-employed "
        "applicants).",
        "Politically Exposed Persons (PEPs) are, by default, categorised as High risk and subjected to "
        "enhanced due diligence (EDD) and senior management approval before onboarding.",
        "No account is opened in anonymous or fictitious names, or where the company is unable to apply "
        "appropriate customer due diligence.",
    ])

    add_heading(doc, "3. Officially Valid Documents (OVDs) Accepted")
    add_table(doc, ["Document", "Used For", "Notes"], [
        ["Aadhaar (masked/last 4 digits displayed)", "Identity + Address proof", "Preferred via Aadhaar-based e-KYC/Digilocker where consented"],
        ["PAN Card", "Identity + tax linkage", "Mandatory for loans above the prescribed threshold"],
        ["Voter ID Card", "Identity + Address proof", "Accepted as an alternate OVD"],
        ["Passport", "Identity + Address proof", "Accepted; NRI applicants require additional documentation"],
        ["Driving Licence", "Identity + Address proof", "Accepted as an alternate OVD"],
        ["NREGA Job Card", "Identity + Address proof", "Accepted for rural/microfinance segment"],
    ])

    add_heading(doc, "4. Periodic Updation (Re-KYC)")
    add_body(doc,
        "Re-KYC is carried out based on risk category: every 2 years for High risk customers, every 8 "
        "years for Medium risk, and every 10 years for Low risk customers, or earlier if there is a "
        "material change in the customer's profile (address, occupation, or unusual transaction pattern).")

    add_heading(doc, "5. Monitoring and Reporting")
    add_bullets(doc, [
        "Transactions inconsistent with a customer's declared income or profile are flagged for review "
        "by the AML/Compliance desk.",
        "Cash transactions above the prescribed threshold, and any suspicious transaction, are reported "
        "to the Financial Intelligence Unit-India (FIU-IND) as a Cash Transaction Report (CTR) / "
        "Suspicious Transaction Report (STR) as applicable.",
        "Sanctions screening (UN Security Council list, and other applicable lists) is run at onboarding "
        "and periodically thereafter for all active borrowers.",
    ])

    add_heading(doc, "6. Record Retention")
    add_body(doc,
        "KYC records, account opening documents, and transaction records are retained for a minimum of "
        "5 years from the date of transaction, or 5 years from the date of cessation of the relationship, "
        "whichever is later, and are made available to regulators/law enforcement on lawful request.")

    return doc


# ---------------------------------------------------------------------
# Document 3: IRAC / NPA Classification and Provisioning Norms
# ---------------------------------------------------------------------

def build_npa_provisioning_norms() -> Document:
    doc = Document()
    add_title(doc, "Income Recognition, Asset Classification and Provisioning (IRACP) Norms", "IRACP-003 v2.1")

    add_heading(doc, "1. Purpose")
    add_body(doc,
        "This document summarises how LendWell classifies loan accounts by days-past-due (DPD), recognises "
        "income, and provisions for expected losses, aligned in spirit with RBI's prudential norms for "
        "NBFCs on income recognition, asset classification and provisioning.")

    add_heading(doc, "2. DPD-Based Classification (Special Mention Accounts and NPA)")
    add_table(doc, ["Bucket", "DPD Range", "Category", "Description"], [
        ["Standard", "0 days", "Standard", "No overdue amount; account performing as per schedule."],
        ["SMA-0", "1-30 days", "Standard (early warning)", "Principal/interest overdue but not yet 31 days."],
        ["SMA-1", "31-60 days", "Standard (early warning)", "Overdue between 31 and 60 days."],
        ["SMA-2", "61-90 days", "Standard (early warning)", "Overdue between 61 and 90 days - last stage before NPA."],
        ["NPA - Sub-Standard", "91-365 days", "Non-Performing", "Overdue beyond 90 days; classified NPA for term loans."],
        ["NPA - Doubtful", "366-730 days", "Non-Performing", "Remained Sub-Standard for 12 months."],
        ["NPA - Loss", "731+ days / identified as loss", "Non-Performing", "Considered uncollectible; full provisioning required."],
    ])

    add_heading(doc, "3. Provisioning Norms Applied")
    add_table(doc, ["Asset Classification", "Provisioning Requirement (illustrative)"], [
        ["Standard (incl. SMA-0/1)", "0.40% of outstanding, generally"],
        ["SMA-2", "Higher provisioning as a precautionary buffer (typically 10-15%)"],
        ["Sub-Standard", "Minimum 25% of the unsecured portion; higher for secured shortfall"],
        ["Doubtful", "25-100% depending on how long the account has been doubtful and security cover"],
        ["Loss", "100% - the entire asset is written off or fully provided for"],
    ])

    add_heading(doc, "4. Income Recognition")
    add_body(doc,
        "Interest income on Standard assets (including SMA) is recognised on an accrual basis. Once an "
        "account is classified NPA, further interest income is recognised only on actual realisation "
        "(cash basis), and any interest already accrued but not received is reversed.")

    add_heading(doc, "5. Upgradation to Standard")
    add_body(doc,
        "An NPA account may be upgraded to Standard only after the borrower clears all arrears of "
        "principal and interest, and demonstrates satisfactory performance during a subsequent monitoring "
        "period, as approved under the credit policy - a single part-payment does not by itself justify "
        "an upgrade.")

    add_heading(doc, "6. Restructuring")
    add_body(doc,
        "Restructured accounts (arising from genuine financial difficulty of the borrower, e.g. a revised "
        "repayment schedule or moratorium) are classified and provisioned as per the applicable "
        "restructuring framework and are separately flagged and reported in the NPA_CLASSIFICATION table "
        "for audit visibility.")

    return doc


# ---------------------------------------------------------------------
# Document 4: Digital Lending Guidelines
# ---------------------------------------------------------------------

def build_digital_lending_guidelines() -> Document:
    doc = Document()
    add_title(doc, "Digital Lending Operating Guidelines", "DLG-005 v1.4")

    add_heading(doc, "1. Scope")
    add_body(doc,
        "Applies to loans sourced or serviced through the company's own mobile/web app, and through "
        "Lending Service Providers (LSPs) - digital platforms and DSAs that source applications on the "
        "company's behalf - consistent with the regulatory direction on digital lending.")

    add_heading(doc, "2. Direct Disbursement and Repayment")
    add_bullets(doc, [
        "Loan disbursement is made directly into the borrower's bank account; no pass-through or pool "
        "account of any third-party LSP is used to route disbursement or repayment.",
        "All fees payable to an LSP are paid directly by the company, and are not collected from the "
        "borrower separately or deducted from the disbursed loan amount.",
    ])

    add_heading(doc, "3. Key Fact Statement (KFS)")
    add_body(doc,
        "Before loan execution, the borrower is provided a standardised Key Fact Statement summarising: "
        "the Annual Percentage Rate (APR), the amount and purpose of every fee/charge, the recovery "
        "mechanism, details of the grievance redressal officer, and the cooling-off/look-up period during "
        "which the borrower may exit the loan by paying only the principal and proportionate APR, without "
        "penalty.")

    add_heading(doc, "4. Data Privacy and Consent")
    add_bullets(doc, [
        "Only data necessary for the credit decision is collected, with explicit borrower consent, and "
        "the borrower can revoke consent for future data collection (subject to already-disbursed loans "
        "being serviced).",
        "Borrower data is not shared with the LSP or any third party beyond what is required for the "
        "specific loan, and is not used for cross-selling without separate, specific consent.",
        "Biometric/Aadhaar data is stored only in masked/encrypted form; full Aadhaar numbers are never "
        "persisted in application databases.",
    ])

    add_heading(doc, "5. Algorithmic Lending Decisions")
    add_body(doc,
        "Where a scorecard/model is used to assist the credit decision, the model's key input features "
        "and the rationale for rejection are auditable, and a customer can request the principal reason "
        "for rejection of an application.")

    add_heading(doc, "6. Grievance Redressal for Digital Loans")
    add_body(doc,
        "A dedicated nodal grievance officer handles digital-lending complaints, with an internal "
        "resolution target of 30 days, escalating to the RBI Ombudsman framework for unresolved "
        "complaints, per the Grievance Redressal Policy (GRP-004).")

    return doc


# ---------------------------------------------------------------------
# Document 5: Grievance Redressal Policy
# ---------------------------------------------------------------------

def build_grievance_redressal_policy() -> Document:
    doc = Document()
    add_title(doc, "Customer Grievance Redressal Policy", "GRP-004 v2.3")

    add_heading(doc, "1. Objective")
    add_body(doc,
        "To provide customers a fair, time-bound, and accessible mechanism to raise and get complaints "
        "resolved regarding any product, service, or conduct of the company or its representatives/agents, "
        "aligned with RBI's Integrated Ombudsman Scheme framework for regulated entities.")

    add_heading(doc, "2. Channels for Raising a Complaint")
    add_bullets(doc, [
        "Toll-free customer care number, available on all loan documents and the company website.",
        "Dedicated email address monitored by the Customer Service Cell.",
        "In person at any branch, or through the mobile app's 'Help & Support' section.",
    ])

    add_heading(doc, "3. Escalation Matrix and Turnaround Time")
    add_table(doc, ["Level", "Escalation Point", "Resolution TAT"], [
        ["1", "Branch Manager / Customer Service Cell", "7 working days"],
        ["2", "Regional Grievance Redressal Officer", "15 working days from original complaint date"],
        ["3", "Principal Nodal Officer (Head Office)", "30 working days from original complaint date"],
        ["4", "RBI Ombudsman (if unresolved after 30 days, or reply unsatisfactory)", "As per Ombudsman Scheme timelines"],
    ])

    add_heading(doc, "4. Specific Provisions for Collections-Related Grievances")
    add_body(doc,
        "Complaints alleging harassment, inappropriate contact hours, or misconduct by a collections "
        "agent are treated as priority and investigated by the Compliance function within 3 working days; "
        "the concerned agent/DSA is suspended from collection activity pending investigation outcome, "
        "per the Fair Practices Code (FPC-001).")

    add_heading(doc, "5. Reporting to the Board")
    add_body(doc,
        "A consolidated grievance report - volume, category, average TAT, and root-cause themes - is "
        "placed before the Customer Service Committee of the Board every quarter.")

    return doc


DOCUMENTS = [
    ("fair_practices_code.docx", build_fair_practices_code),
    ("kyc_aml_policy.docx", build_kyc_aml_policy),
    ("npa_provisioning_norms.docx", build_npa_provisioning_norms),
    ("digital_lending_guidelines.docx", build_digital_lending_guidelines),
    ("grievance_redressal_policy.docx", build_grievance_redressal_policy),
]


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    for filename, builder in DOCUMENTS:
        doc = builder()
        path = os.path.join(OUTPUT_DIR, filename)
        doc.save(path)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
